from tkinter import Button, Label, Entry, Canvas, PhotoImage, Toplevel, Tk, messagebox
from tkinter import filedialog as fd
from tkinter.ttk import LabelFrame
import json
import requests
import audio_effects as ae
from pydub import AudioSegment
from pydub.playback import _play_with_simpleaudio as play_sound
from moviepy.editor import *
import mutagen
from scipy.io import wavfile
from mutagen.wave import WAVE
import moviepy
from moviepy.editor import *
import docx
import threading


def core (filename, Khorugi):
    f = open('text.txt','w' , encoding="utf-8") 
    f.close()    
    API_URL = "https://api-inference.huggingface.co/models/m3hrdadfi/wav2vec2-large-xlsr-persian"
    headers = {"Authorization": f"Bearer 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"}

    def query(filename):
        with open(filename, "rb") as f:
            data = f.read()
        try:
            response = requests.request("POST", API_URL, headers=headers, data=data)
            return json.loads(response.content.decode("utf-8"))
        except :
            pass

    # استخراج صدا از ویدئو
    my_clip = moviepy.editor.VideoFileClip(filename)
    my_clip.audio.write_audiofile("Speech/main.wav")
    current_audio = AudioSegment.from_file("Speech/main.wav")

    # حساب کردن مدت زمان  صدا
    def audio_duration(length):
        hours = length // 3600
        length %= 3600
        mins = length // 60
        length %= 60
        seconds = length
        return hours, mins, seconds

    audio = WAVE("Speech/main.wav") ; audio_info = audio.info ; length = int(audio_info.length)
    hours, mins, seconds = audio_duration(length)
    start = 0 ; end = 10
    All = (hours * (60**2) + mins*60 + seconds)//10
    bagh = (hours * (60**2) + mins*60 + seconds)%10

    # اصل ماجرا |: :/ |:
    def trim_wav( originalWavPath, newWavPath , start, end ):
            sampleRate, waveData = wavfile.read( originalWavPath )
            startSample = int( start * sampleRate )
            endSample = int( end * sampleRate )
            wavfile.write( newWavPath, sampleRate, waveData[startSample:endSample])

    path = "Speech/main.wav"
    TEXT = ''
    keeper = ''

    def Dorost_Kardan_Matn (TEXT):
        TEXT = list(TEXT)

        while "{" in TEXT:
            TEXT.remove("{")

        while "}" in TEXT:
            TEXT.remove("}")

        while ":" in TEXT:
            TEXT.remove(":")

        D = "'text'"

        for x in D:
            while x in TEXT:
                TEXT.remove(x)

        S = ''

        for x in TEXT:
            S += str(x)

        return S

    
        
    for i in range(All+1):
        trim_wav(path, 'Speech/bakh.wav', start , end )
        TEXT = query('Speech/bakh.wav')
        Matn = Dorost_Kardan_Matn(str(TEXT))
        
        if( TEXT != None):
                f = open('text.txt','a' , encoding="utf-8") 
                f.write(str(Matn)) 
                f.close()
        start += 10 ; end += 10


    def Word_Export():
        doc = docx.Document()
        f = open('text.txt' , encoding="utf-8").read()
        doc.add_heading('جزوه ی آماده سازی شده با جــــارو', 0)
        doc_para = doc.add_paragraph(f)
        doc.save(Khorugi + '/Jaroo.docx')

        return 0

    Word_Export()

    print('End of Jaroo :)')


def Sehat ():
    ERRORING = False
    try:
        filename_get = Folder
        filepath_get = File_path.get()

    except:
        ERRORING = True
        messagebox.showerror(title='File Error', message='You need to choose an mp4 file and a Folder to store the result')

    if ERRORING == False:
        if filename_get and filepath_get != "":
            threading.Thread(target=core (filepath_get, filename_get)).start()

        else:
            messagebox.showerror(title='File Error', message='You need to choose an mp4 file and a Folder to store the result')

def frontend_ui ():
    global Folder
    
    def Bast_Dadan_mp4 ():
        global filename
        global File_path
        
        File_path.insert('end',"")
        File_path.insert('end',filename)

    def Bast_Dadan_word ():
        global File_name
        global Folder

        File_name['text'] = Folder    

    def Browse_filee_word():
        global Folder

        Folder = fd.askdirectory(title = 'Open a Folder')

        Bast_Dadan_word ()

    def Browse_filee_mp4():
        global filename

        filetypes = (
                ('Video files', '*.mp4'),
                ('All file types' , '.')

        )

        filename = fd.askopenfilename(
        title='Open a Video',
        initialdir='/',
        filetypes=filetypes)

        Bast_Dadan_mp4 ()

    def about_win ():
        About_win = Toplevel()

        About_win.title('About')
        Label(About_win , image=About_im , bg = 'white').place(x = 0,y = 0)

        About_win.resizable(width = False , height= False)
        About_win.iconphoto(False , Jaroo_icon_im)

        About_win['bg'] = 'white'
        About_win.geometry('290x252')

    def main_win ():
        global canvas
        global File_name
        global File_path

        canvas = Canvas(
            App,
            bg = "white",
            height = 680,
            width = 1080,
            bd = 0,
            highlightthickness = 0,
            relief = "ridge")
        canvas.create_text(
            835, 65,
            text = "Jaroo",
            fill = "Black",
            font = ("Arial", int(32)))
        canvas.create_image(
            280 , 340,
            image = Jaroo_big_im)
        
        canvas.create_image(
            860, 254,
            image = Text_Box_2_im)
        canvas.create_image(
            830, 390,
            image = Text_Box_1_im)

        File_path = Entry(
            App,
            bd = 0,
            bg = "#dcc9b7",
            highlightthickness = 0,)
        File_name = Button(
            App,
            bd = 0,
            bg = "#dcc9b7",
            highlightthickness = 0,
            command=Browse_filee_word
            )  

        Browse_btn = Button(App, image = Browse_im , background='white' , borderwidth=0, cursor='hand2' , command=Browse_filee_mp4)
        Browse_btn.place(x =580 , y = 207)

        canvas.place(x = 0, y = 0)
        File_name.place(x = 630 , y = 376,width = 400,height = 25)
        File_path.place(x = 695, y = 241,width = 300,height = 30)

        About_b = Button(App , image=Dot_im , bg = 'white' ,borderwidth=0 , cursor='hand2' , command=about_win )
        About_b.place(x = 1030 , y = 10)

        Start_b = Button(App , image=Export_Word_im, bg= 'white' , borderwidth=0 , cursor='hand2' , command = Sehat)
        Start_b.place(x = 630 , y = 490)


    main_win ()

App = Tk()
App.title('Jaroo 2.8.3')


Browse_im = PhotoImage(file="assets\Browse_icon.png")
Export_Word_im = PhotoImage(file="assets\Export_Word.png")
Jaroo_big_im = PhotoImage(file="assets\Jaroo_big.png")
Jaroo_icon_im = PhotoImage(file="assets\Jaroo_icon.png")
Jaroo_Name_im = PhotoImage(file="assets\Jaroo_Name.png")
Text_Box_1_im = PhotoImage(file="assets\T2.png")
Text_Box_2_im = PhotoImage(file="assets\TETET.png")
Dot_im = PhotoImage(file='assets\dot.png')
About_im = PhotoImage(file = 'assets\About.png')

App['bg'] = 'white'
App.geometry('1080x680')
App.resizable(width = False , height= False)
App.iconphoto(False , Jaroo_icon_im)

frontend_ui ()
App.mainloop()
